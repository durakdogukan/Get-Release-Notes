from selenium import webdriver
import re
import docx

driver_path = "C:\\Users\\Dogukan\\Desktop\\python\\msedgedriver.exe"
browser = webdriver.Edge(executable_path=driver_path)

browser.get("https://www.thunderbird.net/en-US/thunderbird/releases/")

links = browser.find_elements_by_tag_name('a')

release_links = []

for link in links:
    href_link = link.get_attribute('href')
    if href_link.find("/releasenotes/") != -1:
        versionNo = int(re.findall(r'\d+', href_link)[0])
        if versionNo > 30:
            release_links.append(href_link)

#______________________

mydoc = docx.Document()

index = 0

for link in release_links:
    browser.get(link)
    try:
        c_release_version = browser.find_element_by_xpath('//*[@id="masthead"]/section[2]/p').text

        mydoc.add_heading(c_release_version, 0)

        contents = browser.find_elements_by_xpath('/html/body/main/section[3]/aside')
        for content in contents:
            print(content.text)
            mydoc.add_paragraph(' - ' + content.text)
        print('\n')
    except:
        print('hata')

browser.quit()
mydoc.save("C:\\Users\\Dogukan\\Desktop\\python\\Release\\ThunderbirdReleases.docx")
