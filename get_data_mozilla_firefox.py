from selenium import webdriver
import re
import docx

driver_path = "C:\\Users\\Dogukan\\Desktop\\python\\msedgedriver.exe"
browser = webdriver.Edge(executable_path=driver_path)

browser.get("https://www.mozilla.org/en-US/firefox/releases/")

links = browser.find_elements_by_tag_name('a')

release_links = []

# for link in links:
#     href_link = link.get_attribute('href')
#     versionNo = float(re.findall(r'\d+', str(href_link))[0])
#     if href_link.find("releasenotes") != -1 and versionNo >= 28.0:
#         release_links.append(href_link)


for link in links:
    href_link = link.get_attribute('href')
    if href_link.find("releasenotes") != -1:
        versionNo = int(re.findall(r'\d+', href_link)[0])
        if versionNo > 28:
            release_links.append(href_link)        

#______________________

mydoc = docx.Document()

index = 0

for link in release_links:
    browser.get(link)
    try:
        c_release_version = browser.find_element_by_class_name('c-release-version').text
        c_release_date = browser.find_element_by_class_name('c-release-date').text

        mydoc.add_heading(c_release_version + ' ' + c_release_date, 0)

        contents = browser.find_elements_by_xpath('//*[@id="main-content"]/section[1]/div')
        for content in contents:
            print('*', content.text)
            mydoc.add_paragraph(' - ' + content.text)
        print('\n')
    except:
        print('hata')

browser.quit()
mydoc.save("C:\\Users\\Dogukan\\Desktop\\python\\Release\\MozillaFirefoxReleasesx.docx")